"""
docx/xlsx → XML展開 & テキスト抽出スクリプト
Open XML形式のファイル（docx/xlsx/xlsm/pptx）をZIP解凍してXMLを取り出し、
docxの場合はテキスト抽出も行う。

使い方:
    python extract_docx_xml.py [--src フォルダパス]

オプション:
    --src  対象フォルダ（デフォルト: カレントディレクトリ）

前提:
    - pip install python-docx
"""

import glob
import os
import zipfile
import argparse
from docx import Document


def extract_xml(file_path: str, out_dir: str):
	"""Open XMLファイルをZIP解凍してXMLを展開"""
	os.makedirs(out_dir, exist_ok=True)
	with zipfile.ZipFile(file_path, 'r') as z:
		z.extractall(out_dir)
	return out_dir


def extract_text(docx_path: str) -> str:
	"""docxからプレーンテキストを抽出"""
	doc = Document(docx_path)
	lines = []
	for para in doc.paragraphs:
		lines.append(para.text)
	# テーブルも抽出
	for table in doc.tables:
		for row in table.rows:
			cells = [cell.text.strip() for cell in row.cells]
			lines.append(" | ".join(cells))
	return "\n".join(lines)


def main():
	parser = argparse.ArgumentParser(description="docx/xlsx → XML展開 & テキスト抽出")
	parser.add_argument("--src", default=".", help="対象フォルダ")
	args = parser.parse_args()

	src_dir = os.path.abspath(args.src)
	xml_base = os.path.join(src_dir, "_xml")
	txt_base = os.path.join(src_dir, "_txt")

	# 対象拡張子
	target_exts = ("*.docx", "*.xlsx", "*.xlsm", "*.pptx")
	all_files = []
	for ext in target_exts:
		all_files.extend(glob.glob(os.path.join(src_dir, "**", ext), recursive=True))

	# _xml, _txt フォルダ内のファイルは除外、~$ 一時ファイルも除外
	all_files = [f for f in all_files
				 if "_xml" not in f and "_txt" not in f
				 and not os.path.basename(f).startswith("~$")]

	if not all_files:
		print("変換対象のファイルなし")
		return

	print(f"=== Open XML展開 & テキスト抽出 ===")
	print(f"対象: {len(all_files)} 件")
	print(f"XML出力先: {xml_base}")
	print(f"テキスト出力先: {txt_base}")
	print()

	for file_path in all_files:
		name = os.path.splitext(os.path.basename(file_path))[0]
		ext = os.path.splitext(file_path)[1].lower()
		print(f"[処理中] {os.path.basename(file_path)}")

		# XML展開
		xml_out = os.path.join(xml_base, name)
		try:
			extract_xml(file_path, xml_out)
			print(f"  [XML]  → {os.path.relpath(xml_out, src_dir)}")
		except Exception as e:
			print(f"  [XML失敗] {e}")

		# テキスト抽出（docxのみ）
		if ext == ".docx":
			os.makedirs(txt_base, exist_ok=True)
			txt_out = os.path.join(txt_base, name + ".txt")
			try:
				text = extract_text(file_path)
				with open(txt_out, "w", encoding="utf-8") as f:
					f.write(text)
				print(f"  [TXT]  → {os.path.relpath(txt_out, src_dir)}")
			except Exception as e:
				print(f"  [TXT失敗] {e}")

	print("\n=== 完了 ===")


if __name__ == "__main__":
	main()
